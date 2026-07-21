import docx
import pytest

from loader.docx.python_docx import Loader
from util.passthrough_layout import PassthroughLayout

pytestmark = pytest.mark.unit


@pytest.fixture
def docx_file(tmp_path):
    document = docx.Document()
    document.add_paragraph("Intro paragraph", style="Normal")
    document.add_heading("Section One", level=1)
    document.add_paragraph("Body text", style="Normal")
    document.add_paragraph("Bullet one", style="List Bullet")
    document.add_paragraph("Bullet two", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "H1"
    table.rows[0].cells[1].text = "H2"
    table.rows[1].cells[0].text = "A"
    table.rows[1].cells[1].text = "B"
    document.add_paragraph("Outro paragraph", style="Normal")

    path = tmp_path / "sample.docx"
    document.save(path)
    return str(path)


def test_get_layout_and_ocr(docx_file):
    loader = Loader(layout_config={"type": "detr"}, ocr_config={"type": "paddle"})
    assert isinstance(loader.layout, PassthroughLayout)
    assert loader.ocr is None


def test_preserves_document_order_and_assigns_categories(docx_file):
    loader = Loader()
    items = loader(docx_file)

    categories = [(item["text"], item["category"]) for item in items]
    assert categories[0] == ("Intro paragraph", "text")
    assert categories[1] == ("Section One", "section_header")
    assert categories[2] == ("Body text", "text")
    assert categories[3] == ("Bullet one", "list_item")
    assert categories[4] == ("Bullet two", "list_item")
    assert categories[5][1] == "table"
    assert "<td>H1</td>" in categories[5][0]
    assert "<td>A</td>" in categories[5][0]
    assert categories[6] == ("Outro paragraph", "text")
    assert all(item["page"] == 1 for item in items)


def test_furniture_wrapper_table_is_unwrapped(sample_dir):
    loader = Loader()
    items = loader(str(sample_dir / "docx" / "tablecell.docx"))

    categories = [item["category"] for item in items]
    texts = [item["text"] for item in items]

    assert categories.count("table") == 2
    assert "Some text before" in texts
    assert "Some text after" in texts
    table_texts = [item["text"] for item in items if item["category"] == "table"]
    assert any("Tab1" in text for text in table_texts)
